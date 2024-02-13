import os
from typing import List
import docx
import chromadb
import tabulate
from .embedding import SentencesToEmbeddings


def extract_tables(doc):
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text
                row_data.append(cell_text)
            table_data.append(row_data)

        tables.append(tabulate(table_data, tablefmt="plain"))
    return tables


def extract_texts(doc):
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    return texts


def consolidate_elements(elements, sentence_size=256, overlapping_num=3):

    final_sentence_list = []
    accumulate_len = 0
    windows_sentences = []

    for sentence in elements:

        word_len = len(sentence)
        if accumulate_len+word_len <= sentence_size:
            windows_sentences.append(sentence)
            accumulate_len += word_len
        else:
            windows_sentence = "\n".join(windows_sentences)
            final_sentence_list.append(windows_sentence)
            windows_sentences = windows_sentences[-overlapping_num:].copy()+[
                sentence]
            accumulate_len = word_len

    if len(windows_sentences) > 0:
        windows_sentence = "\n".join(windows_sentences)
        final_sentence_list.append(windows_sentence)

    return final_sentence_list


class ChromaEngine():
    def __init__(self, db_path="/Users/nelsonlin/Desktop/workspaces/llm-local-chat-api/app/data") -> None:
        self.chroma_client = chromadb.PersistentClient(path=db_path)
        self.document_collection = self.chroma_client.get_or_create_collection(
            name="documents")
        self.embedding_collection = self.chroma_client.get_or_create_collection(
            name="documents_embedding")
        self.sentences_to_embeddings = SentencesToEmbeddings()

    def ingest_file(self, file_path, sentence_size=256, overlapping_num=2, overwrite=True):

        file_path = os.path.abspath(file_path)
        file_name = os.path.basename(file_path)

        if overwrite:
            self.document_collection.delete(where={"file_name": file_name})

        self.document_collection.add(
            documents=[file_name],
            embeddings=[[0]],
            metadatas={"file_name": file_name, "file_path": file_path},
            ids=[file_name]
        )

        chunks = self.parse_docx(
            file_path, sentence_size=sentence_size,
            overlapping_num=overlapping_num)

        embeddings = self.sentences_to_embeddings(chunks)
        ids = [file_name+f"_{i}" for i in range(len(chunks))]

        metadatas = [{'file_path': file_path,
                      'file_name': file_name}] * len(chunks)

        if overwrite:
            self.embedding_collection.delete(where={"file_name": file_name})

        self.embedding_collection.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids
        )

    def parse_docx(self, file_path, sentence_size=256, overlapping_num=2) -> List[str]:
        doc = docx.Document(file_path)
        tables = extract_tables(doc)

        elements = []
        table_indices = []
        for index, element in enumerate(doc.element.body):
            if element.tag.endswith("tbl"):
                table_indices.append(index)
            elements.append(element.text)

        for index, table in enumerate(tables):
            table_index = table_indices[index]
            elements[table_index] = table

        elements = [
            element for element in elements if isinstance(element, str)]

        chunks = consolidate_elements(
            elements,
            sentence_size=sentence_size,
            overlapping_num=overlapping_num)

        return chunks

    def list_document(self):
        return self.document_collection.get()

    def vector_search(self, file_name: str, query_text: str, limit=5) -> List[str]:

        query_embedding = self.sentences_to_embeddings([query_text])[0]

        embedding_query_results = self.embedding_collection.query(query_embeddings=query_embedding,
                                                                  where={
                                                                      "file_name": file_name},
                                                                  n_results=limit)

        results = embedding_query_results['documents'][0]
        return "\n".join(results)
